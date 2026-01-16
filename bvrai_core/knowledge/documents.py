"""
Document Processing Module

This module handles document loading and text extraction
from various file formats.
"""

import hashlib
import logging
import mimetypes
import os
import re
import time
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import (
    Any,
    BinaryIO,
    Dict,
    List,
    Optional,
    Type,
    Union,
)

from .base import (
    DocumentMetadata,
    DocumentType,
    ChunkMetadata,
    detect_document_type,
    generate_document_id,
    generate_chunk_id,
    estimate_tokens,
)


logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Represents a loaded document."""

    id: str
    content: str
    metadata: DocumentMetadata

    # Raw content (before cleaning)
    raw_content: Optional[str] = None

    # Structure information
    sections: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)

    # Content hash for deduplication
    content_hash: str = ""

    def __post_init__(self):
        if not self.content_hash and self.content:
            self.content_hash = hashlib.sha256(self.content.encode()).hexdigest()

        # Update metadata counts
        self.metadata.char_count = len(self.content)
        self.metadata.word_count = len(self.content.split())

    @property
    def estimated_tokens(self) -> int:
        """Estimate token count."""
        return estimate_tokens(self.content)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "id": self.id,
            "content": self.content,
            "metadata": self.metadata.to_dict(),
            "content_hash": self.content_hash,
            "sections": self.sections,
        }


@dataclass
class DocumentChunk:
    """Represents a chunk of a document."""

    id: str
    content: str
    metadata: ChunkMetadata
    embedding: Optional[List[float]] = None

    # Parent document reference
    document_id: str = ""

    def __post_init__(self):
        self.metadata.char_count = len(self.content)
        self.metadata.word_count = len(self.content.split())
        self.metadata.token_count = estimate_tokens(self.content)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "id": self.id,
            "content": self.content,
            "metadata": self.metadata.to_dict(),
            "document_id": self.document_id,
        }


class DocumentProcessor(ABC):
    """Abstract base class for document processors."""

    supported_types: List[DocumentType] = []

    @abstractmethod
    def process(
        self,
        source: Union[str, Path, BinaryIO],
        metadata: Optional[DocumentMetadata] = None,
    ) -> Document:
        """Process a document and extract text."""
        pass

    def clean_text(self, text: str) -> str:
        """Clean extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)

        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)

        return text.strip()


class TextProcessor(DocumentProcessor):
    """Processor for plain text files."""

    supported_types = [DocumentType.TXT]

    def process(
        self,
        source: Union[str, Path, BinaryIO],
        metadata: Optional[DocumentMetadata] = None,
    ) -> Document:
        """Process a text file."""
        # Read content
        if isinstance(source, (str, Path)):
            path = Path(source)
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            source_str = str(path)
            file_size = path.stat().st_size
        else:
            content = source.read()
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='replace')
            source_str = getattr(source, 'name', 'stream')
            file_size = len(content)

        # Create metadata
        if metadata is None:
            metadata = DocumentMetadata(
                source=source_str,
                source_type=DocumentType.TXT,
                file_size_bytes=file_size,
            )

        # Clean content
        clean_content = self.clean_text(content)

        # Generate ID
        doc_id = generate_document_id(source_str)

        return Document(
            id=doc_id,
            content=clean_content,
            raw_content=content,
            metadata=metadata,
        )


class PDFProcessor(DocumentProcessor):
    """Processor for PDF files."""

    supported_types = [DocumentType.PDF]

    def process(
        self,
        source: Union[str, Path, BinaryIO],
        metadata: Optional[DocumentMetadata] = None,
    ) -> Document:
        """Process a PDF file."""
        try:
            import pypdf
        except ImportError:
            raise ImportError("pypdf package required. Install with: pip install pypdf")

        # Read PDF
        if isinstance(source, (str, Path)):
            path = Path(source)
            reader = pypdf.PdfReader(str(path))
            source_str = str(path)
            file_size = path.stat().st_size
        else:
            reader = pypdf.PdfReader(source)
            source_str = getattr(source, 'name', 'stream')
            file_size = 0

        # Extract text from all pages
        pages_content = []
        sections = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages_content.append(text)
            sections.append({
                "type": "page",
                "page_number": i + 1,
                "content": text,
            })

        content = "\n\n".join(pages_content)

        # Extract metadata from PDF
        pdf_metadata = reader.metadata or {}

        if metadata is None:
            metadata = DocumentMetadata(
                source=source_str,
                source_type=DocumentType.PDF,
                file_size_bytes=file_size,
                page_count=len(reader.pages),
                title=pdf_metadata.get('/Title', ''),
                author=pdf_metadata.get('/Author', ''),
            )
        else:
            metadata.page_count = len(reader.pages)

        # Clean content
        clean_content = self.clean_text(content)

        # Generate ID
        doc_id = generate_document_id(source_str)

        return Document(
            id=doc_id,
            content=clean_content,
            raw_content=content,
            metadata=metadata,
            sections=sections,
        )


class DocxProcessor(DocumentProcessor):
    """Processor for Word documents."""

    supported_types = [DocumentType.DOCX, DocumentType.DOC]

    def process(
        self,
        source: Union[str, Path, BinaryIO],
        metadata: Optional[DocumentMetadata] = None,
    ) -> Document:
        """Process a Word document."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx package required. Install with: pip install python-docx")

        # Read document
        if isinstance(source, (str, Path)):
            path = Path(source)
            doc = DocxDocument(str(path))
            source_str = str(path)
            file_size = path.stat().st_size
        else:
            doc = DocxDocument(source)
            source_str = getattr(source, 'name', 'stream')
            file_size = 0

        # Extract text from paragraphs
        paragraphs = []
        sections = []
        current_section = {"type": "body", "content": []}

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

                # Detect headings
                if para.style and 'Heading' in para.style.name:
                    level = 1
                    try:
                        level = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        pass

                    if current_section["content"]:
                        sections.append(current_section)

                    current_section = {
                        "type": "section",
                        "title": text,
                        "level": level,
                        "content": [],
                    }
                else:
                    current_section["content"].append(text)

        if current_section["content"]:
            sections.append(current_section)

        content = "\n\n".join(paragraphs)

        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append({"data": table_data})

        # Create metadata
        if metadata is None:
            core_props = doc.core_properties
            metadata = DocumentMetadata(
                source=source_str,
                source_type=DocumentType.DOCX,
                file_size_bytes=file_size,
                title=core_props.title or "",
                author=core_props.author or "",
            )

        # Clean content
        clean_content = self.clean_text(content)

        # Generate ID
        doc_id = generate_document_id(source_str)

        return Document(
            id=doc_id,
            content=clean_content,
            raw_content=content,
            metadata=metadata,
            sections=sections,
            tables=tables,
        )


class HTMLProcessor(DocumentProcessor):
    """Processor for HTML content."""

    supported_types = [DocumentType.HTML]

    def process(
        self,
        source: Union[str, Path, BinaryIO],
        metadata: Optional[DocumentMetadata] = None,
    ) -> Document:
        """Process HTML content."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 package required. Install with: pip install beautifulsoup4")

        # Read content
        if isinstance(source, (str, Path)):
            path = Path(source)
            if path.exists():
                with open(path, 'r', encoding='utf-8', errors='replace') as f:
                    html_content = f.read()
                source_str = str(path)
                file_size = path.stat().st_size
            else:
                # Treat as raw HTML string
                html_content = str(source)
                source_str = "html_string"
                file_size = len(html_content)
        else:
            html_content = source.read()
            if isinstance(html_content, bytes):
                html_content = html_content.decode('utf-8', errors='replace')
            source_str = getattr(source, 'name', 'stream')
            file_size = len(html_content)

        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')

        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header']):
            element.decompose()

        # Extract title
        title = ""
        title_tag = soup.find('title')
        if title_tag:
            title = title_tag.get_text().strip()

        # Extract main content
        main_content = soup.find('main') or soup.find('article') or soup.find('body') or soup

        # Extract text
        text = main_content.get_text(separator='\n')

        # Extract sections from headers
        sections = []
        for header in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            level = int(header.name[1])
            sections.append({
                "type": "heading",
                "title": header.get_text().strip(),
                "level": level,
            })

        # Create metadata
        if metadata is None:
            metadata = DocumentMetadata(
                source=source_str,
                source_type=DocumentType.HTML,
                file_size_bytes=file_size,
                title=title,
            )

        # Clean content
        clean_content = self.clean_text(text)

        # Generate ID
        doc_id = generate_document_id(source_str)

        return Document(
            id=doc_id,
            content=clean_content,
            raw_content=html_content,
            metadata=metadata,
            sections=sections,
        )


class MarkdownProcessor(DocumentProcessor):
    """Processor for Markdown files."""

    supported_types = [DocumentType.MD]

    def process(
        self,
        source: Union[str, Path, BinaryIO],
        metadata: Optional[DocumentMetadata] = None,
    ) -> Document:
        """Process a Markdown file."""
        # Read content
        if isinstance(source, (str, Path)):
            path = Path(source)
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            source_str = str(path)
            file_size = path.stat().st_size
        else:
            content = source.read()
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='replace')
            source_str = getattr(source, 'name', 'stream')
            file_size = len(content)

        # Extract sections from headers
        sections = []
        header_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)

        for match in header_pattern.finditer(content):
            level = len(match.group(1))
            title = match.group(2).strip()
            sections.append({
                "type": "heading",
                "title": title,
                "level": level,
                "position": match.start(),
            })

        # Extract title from first H1 or front matter
        title = ""
        if sections and sections[0]["level"] == 1:
            title = sections[0]["title"]

        # Check for YAML front matter
        front_matter_pattern = re.compile(r'^---\s*\n(.*?)\n---\s*\n', re.DOTALL)
        front_matter_match = front_matter_pattern.match(content)

        if front_matter_match:
            try:
                import yaml
                front_matter = yaml.safe_load(front_matter_match.group(1))
                if isinstance(front_matter, dict):
                    title = front_matter.get('title', title)
            except Exception:
                pass

            # Remove front matter from content
            content = content[front_matter_match.end():]

        # Create metadata
        if metadata is None:
            metadata = DocumentMetadata(
                source=source_str,
                source_type=DocumentType.MD,
                file_size_bytes=file_size,
                title=title,
            )

        # Keep markdown mostly as-is but clean excessive whitespace
        clean_content = self.clean_text(content)

        # Generate ID
        doc_id = generate_document_id(source_str)

        return Document(
            id=doc_id,
            content=clean_content,
            raw_content=content,
            metadata=metadata,
            sections=sections,
        )


class CSVProcessor(DocumentProcessor):
    """Processor for CSV files."""

    supported_types = [DocumentType.CSV]

    def process(
        self,
        source: Union[str, Path, BinaryIO],
        metadata: Optional[DocumentMetadata] = None,
    ) -> Document:
        """Process a CSV file."""
        import csv
        from io import StringIO

        # Read content
        if isinstance(source, (str, Path)):
            path = Path(source)
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            source_str = str(path)
            file_size = path.stat().st_size
        else:
            content = source.read()
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='replace')
            source_str = getattr(source, 'name', 'stream')
            file_size = len(content)

        # Parse CSV
        reader = csv.reader(StringIO(content))
        rows = list(reader)

        # Create text representation
        text_parts = []
        tables = []

        if rows:
            headers = rows[0]
            data_rows = rows[1:]

            # Create readable text from rows
            for row in data_rows[:100]:  # Limit for very large CSVs
                row_text = ", ".join(
                    f"{headers[i]}: {cell}" if i < len(headers) else cell
                    for i, cell in enumerate(row) if cell
                )
                if row_text:
                    text_parts.append(row_text)

            tables.append({
                "headers": headers,
                "row_count": len(data_rows),
                "sample_rows": data_rows[:10],
            })

        text = "\n".join(text_parts)

        # Create metadata
        if metadata is None:
            metadata = DocumentMetadata(
                source=source_str,
                source_type=DocumentType.CSV,
                file_size_bytes=file_size,
            )

        # Generate ID
        doc_id = generate_document_id(source_str)

        return Document(
            id=doc_id,
            content=text,
            raw_content=content,
            metadata=metadata,
            tables=tables,
        )


class JSONProcessor(DocumentProcessor):
    """Processor for JSON files."""

    supported_types = [DocumentType.JSON]

    def process(
        self,
        source: Union[str, Path, BinaryIO],
        metadata: Optional[DocumentMetadata] = None,
    ) -> Document:
        """Process a JSON file."""
        import json

        # Read content
        if isinstance(source, (str, Path)):
            path = Path(source)
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            source_str = str(path)
            file_size = path.stat().st_size
            raw_content = json.dumps(data, indent=2)
        else:
            content = source.read()
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            data = json.loads(content)
            source_str = getattr(source, 'name', 'stream')
            file_size = len(content)
            raw_content = content

        # Convert JSON to readable text
        text = self._json_to_text(data)

        # Create metadata
        if metadata is None:
            metadata = DocumentMetadata(
                source=source_str,
                source_type=DocumentType.JSON,
                file_size_bytes=file_size,
            )

        # Generate ID
        doc_id = generate_document_id(source_str)

        return Document(
            id=doc_id,
            content=text,
            raw_content=raw_content,
            metadata=metadata,
        )

    def _json_to_text(self, data: Any, prefix: str = "") -> str:
        """Convert JSON data to readable text."""
        lines = []

        if isinstance(data, dict):
            for key, value in data.items():
                new_prefix = f"{prefix}{key}"
                if isinstance(value, (dict, list)):
                    lines.append(f"{new_prefix}:")
                    lines.append(self._json_to_text(value, f"  {prefix}"))
                else:
                    lines.append(f"{new_prefix}: {value}")

        elif isinstance(data, list):
            for i, item in enumerate(data[:50]):  # Limit for large arrays
                if isinstance(item, (dict, list)):
                    lines.append(f"{prefix}[{i}]:")
                    lines.append(self._json_to_text(item, f"  {prefix}"))
                else:
                    lines.append(f"{prefix}- {item}")

        else:
            lines.append(f"{prefix}{data}")

        return "\n".join(lines)


# Processor registry
_PROCESSORS: Dict[DocumentType, Type[DocumentProcessor]] = {
    DocumentType.TXT: TextProcessor,
    DocumentType.PDF: PDFProcessor,
    DocumentType.DOCX: DocxProcessor,
    DocumentType.DOC: DocxProcessor,
    DocumentType.HTML: HTMLProcessor,
    DocumentType.MD: MarkdownProcessor,
    DocumentType.CSV: CSVProcessor,
    DocumentType.JSON: JSONProcessor,
}


def get_processor(doc_type: DocumentType) -> Optional[DocumentProcessor]:
    """Get a processor for a document type."""
    processor_class = _PROCESSORS.get(doc_type)
    if processor_class:
        return processor_class()
    return None


def process_document(
    source: Union[str, Path, BinaryIO],
    doc_type: Optional[DocumentType] = None,
    metadata: Optional[DocumentMetadata] = None,
) -> Document:
    """Process a document with automatic type detection."""
    # Detect type if not provided
    if doc_type is None:
        if isinstance(source, (str, Path)):
            doc_type = detect_document_type(str(source))
        else:
            name = getattr(source, 'name', '')
            doc_type = detect_document_type(name)

    # Get processor
    processor = get_processor(doc_type)
    if processor is None:
        # Fall back to text processor
        processor = TextProcessor()

    return processor.process(source, metadata)


__all__ = [
    "Document",
    "DocumentChunk",
    "DocumentProcessor",
    "TextProcessor",
    "PDFProcessor",
    "DocxProcessor",
    "HTMLProcessor",
    "MarkdownProcessor",
    "CSVProcessor",
    "JSONProcessor",
    "get_processor",
    "process_document",
]
